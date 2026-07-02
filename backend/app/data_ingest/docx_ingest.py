# data_ingest/docx_ingest.py
# Parses Word .docx documents using the python-docx library.
# Extracts paragraphs and tables.

import logging
from typing import List

logger = logging.getLogger(__name__)

def load_documents_from_docx(file_path: str) -> List[str]:
    """
    Extracts text paragraphs and table cells from a DOCX file.
    """
    paragraphs = []
    try:
        import docx
        doc = docx.Document(file_path)
        
        # 1. Extract paragraphs text
        for p in doc.paragraphs:
            text = p.text.strip()
            if text:
                paragraphs.append(text)
                
        # 2. Extract tables text
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    paragraphs.append(" | ".join(row_text))
                    
        logger.info(f"Extracted {len(paragraphs)} paragraphs/text units from Word doc: {file_path}")
    except Exception as exc:
        logger.error(f"Error parsing Word file {file_path}: {exc}")
        raise
        
    return paragraphs
